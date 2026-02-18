# /backend/app/services/docx_service.py

from docx import Document

class DocxService:
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from Word document"""
        try:
            doc = Document(file_path)
            full_text = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)

            return "\n".join(full_text)

        except Exception as e:
            print(f"DOCX extraction error: {e}")
            return ""

docx_service = DocxService()